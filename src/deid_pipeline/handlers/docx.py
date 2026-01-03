from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Optional

from .base import ExtractedDocument, SegmentSpec, build_document_from_segments
from ..core.contracts import replacement_key


class DocxHandler:
    extensions = [".docx"]

    def extract(self, input_path: Path, *, language: str) -> ExtractedDocument:
        try:
            from docx import Document  # type: ignore
        except Exception as exc:
            raise ImportError("python-docx is required to extract DOCX files") from exc

        doc = Document(str(input_path))
        segments: list[SegmentSpec] = []

        for paragraph in doc.paragraphs:
            text = str(paragraph.text or "")
            segments.append(SegmentSpec(text=text, metadata={"kind": "paragraph"}))

        return build_document_from_segments(
            input_path=input_path,
            language=language,
            segments=segments or [SegmentSpec(text="")],
            separator="\n",
            metadata={"format": "docx"},
        )

    def rebuild(
        self,
        document: ExtractedDocument,
        *,
        output_text: str,
        entities: list[dict],
        replacement_map: Dict[str, str],
        events: list[dict],
        output_dir: Optional[Path] = None,
        mode: str = "replace",
    ) -> Dict[str, Any]:
        artifacts: Dict[str, Any] = {"output_text": output_text}

        if output_dir is None or mode != "replace":
            artifacts["rebuild_supported"] = False
            return artifacts

        try:
            from docx import Document  # type: ignore
        except Exception:
            artifacts["rebuild_supported"] = False
            return artifacts

        output_dir.mkdir(parents=True, exist_ok=True)
        out_path = output_dir / f"{document.input_path.stem}.deid.docx"

        doc = Document(str(document.input_path))

        # Best-effort: apply replacements to paragraph/table text.
        replacements: list[tuple[str, str]] = []
        for entity in entities:
            original = entity.get("text")
            entity_type = entity.get("type")
            if not original or not entity_type:
                continue
            repl = replacement_map.get(replacement_key(str(entity_type), str(original)))
            if repl is None:
                continue
            replacements.append((str(original), str(repl)))

        for paragraph in doc.paragraphs:
            text = paragraph.text or ""
            for original, repl in replacements:
                text = text.replace(original, repl)
            if text != paragraph.text:
                paragraph.text = text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text or ""
                    new_text = cell_text
                    for original, repl in replacements:
                        new_text = new_text.replace(original, repl)
                    if new_text != cell_text:
                        cell.text = new_text

        doc.save(str(out_path))

        artifacts["output_path"] = str(out_path)
        artifacts["rebuild_supported"] = True
        artifacts["rebuild_mode"] = "docx_text_replace"
        return artifacts
